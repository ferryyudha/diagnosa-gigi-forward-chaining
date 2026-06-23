import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:left w:val="none"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:right w:val="none"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)

def create_document():
    doc = docx.Document()

    # Set page margins (Standard Skripsi: Top 3cm, Bottom 3cm, Left 4cm, Right 3cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.18)      # ~3cm
        section.bottom_margin = Inches(1.18)   # ~3cm
        section.left_margin = Inches(1.57)     # ~4cm
        section.right_margin = Inches(1.18)    # ~3cm

    # Style defaults
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Heading 1 (BAB) Style
    h1_style = doc.styles.add_style('Skripsi Heading 1', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    h1_font = h1_style.font
    h1_font.name = 'Times New Roman'
    h1_font.size = Pt(14)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 0, 0)

    # Heading 2 (Sub-BAB) Style
    h2_style = doc.styles.add_style('Skripsi Heading 2', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    h2_font = h2_style.font
    h2_font.name = 'Times New Roman'
    h2_font.size = Pt(12)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 0, 0)

    # -------------------------------------------------------------
    # BAB IV HASIL DAN PEMBAHASAN
    # -------------------------------------------------------------
    p_bab = doc.add_paragraph('BAB IV', style='Skripsi Heading 1')
    p_bab.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bab.paragraph_format.space_before = Pt(12)
    p_bab.paragraph_format.space_after = Pt(6)

    p_title = doc.add_paragraph('HASIL DAN PEMBAHASAN', style='Skripsi Heading 1')
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(24)

    # --- SECTION A ---
    p_suba = doc.add_paragraph('A. Definisi Masalah dan Penyelesaian', style='Skripsi Heading 2')
    p_suba.paragraph_format.space_before = Pt(12)
    p_suba.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "Penelitian ini ditujukan untuk memecahkan kendala diagnosis pelayanan kesehatan gigi di Praktik Mandiri "
        "Drg. Hj. Rini Sutarti. Berdasarkan rumusan masalah yang telah ditetapkan, terdapat empat aspek analisis utama "
        "yang diteliti, yaitu perancangan sistem pakar berbasis web, penerapan metode Forward Chaining dalam proses "
        "penalaran gejala, visualisasi alur penelusuran logika gejala, serta efektivitas sistem pakar dalam menyajikan "
        "informasi screening diagnosa awal bagi pengguna."
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(12)
    p.add_run(
        "Melalui perancangan aplikasi Diagnosa Penyakit Gigi Metode Forward Chaining, data klinis berupa 25 gejala gigi dan mulut "
        "yang bersumber dari pakar (dokter gigi) dihubungkan secara dinamis menggunakan model penalaran Forward Chaining. "
        "Sistem ini dirancang untuk menjawab seluruh rumusan masalah penelitian guna menghasilkan alat bantu diagnosa awal "
        "yang valid, transparan, dan mudah digunakan oleh masyarakat luas."
    )

    # 1. Definisi Masalah
    p_suba1 = doc.add_paragraph('1. Definisi Masalah (Berdasarkan Rumusan Masalah)', style='Skripsi Heading 2')
    p_suba1.paragraph_format.left_indent = Inches(0.2)
    p_suba1.paragraph_format.space_before = Pt(6)
    p_suba1.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0.4)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "Secara terperinci, definisi masalah yang dihadapi dalam penelitian ini diuraikan berdasarkan empat poin "
        "rumusan masalah utama, yaitu:"
    )

    problems = [
        ("Bagaimana merancang dan membangun sistem pakar penyakit gigi (Rumusan Masalah 1): ", "Adanya kebutuhan mendesak untuk merancang sebuah infrastruktur sistem pakar berbasis web yang handal, aman, dan mudah diakses oleh pasien di Praktik Mandiri Drg. Hj. Rini Sutarti secara real-time, guna mengatasi keterbatasan fisik berupa panjangnya antrean konsultasi dasar di klinik."),
        ("Bagaimana menerapkan metode Forward Chaining (Rumusan Masalah 2): ", "Bagaimana menstrukturkan basis pengetahuan (knowledge base) berupa 8 penyakit dan 25 gejala klinis gigi ke dalam 40 aturan logika IF-THEN, serta mengimplementasikan mesin penalaran Forward Chaining dalam kode PHP untuk menarik kesimpulan penyakit berdasarkan gejala yang dipilih oleh pengguna."),
        ("Bagaimana proses penelusuran gejala hingga menghasilkan diagnosis (Rumusan Masalah 3): ", "Menghilangkan kendala sistem yang bersifat 'black-box' (menyajikan diagnosa tanpa penjelasan). Sistem harus mampu mendokumentasikan dan menyajikan proses penelusuran logika (chaining trace) yang transparan mengenai bagaimana gejala yang dipilih dicocokkan dengan aturan penyakit gigi, lengkap dengan informasi persentase kecocokannya."),
        ("Bagaimana sistem pakar memberikan informasi awal yang membantu pengguna (Rumusan Masalah 4): ", "Bagaimana sistem pakar dapat menyajikan informasi medis awal yang akurat, mencakup deskripsi penyakit dan solusi penanganan darurat yang valid, serta menyediakan layout hasil diagnosa yang ramah printer agar dapat digunakan pasien sebagai rujukan awal saat berobat ke dokter gigi.")
    ]

    for title, desc in problems:
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.6)
        p.paragraph_format.space_after = Pt(4)
        run_title = p.add_run(title)
        run_title.bold = True
        p.add_run(desc)

    # 2. Solusi Penyelesaian Masalah
    p_suba2 = doc.add_paragraph('2. Solusi Penyelesaian Masalah', style='Skripsi Heading 2')
    p_suba2.paragraph_format.left_indent = Inches(0.2)
    p_suba2.paragraph_format.space_before = Pt(12)
    p_suba2.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0.4)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "Untuk menjawab keempat rumusan masalah tersebut, solusi penyelesaian masalah diwujudkan melalui perancangan "
        "dan pembangunan aplikasi Diagnosa Penyakit Gigi Metode Forward Chaining dengan spesifikasi teknis dan fungsional sebagai berikut:"
    )

    solutions = [
        ("Perancangan dan Pembangunan Aplikasi Diagnosa Penyakit Gigi Metode Forward Chaining (Solusi Rumusan Masalah 1): ", "Membangun sistem pakar berbasis web menggunakan bahasa pemrograman PHP dan database MySQL dengan arsitektur client-server. Antarmuka aplikasi dibuat responsif dan menggunakan visual modern premium dark theme dengan struktur tabel relasional yang solid (users, penyakit, gejala, aturan, konsultasi, dan konsultasi_gejala)."),
        ("Penerapan Engine Forward Chaining (Solusi Rumusan Masalah 2): ", "Mengimplementasikan file engine forward_chaining.php yang memuat class 'ForwardChaining'. Ketika pengguna memilih gejala, engine membandingkan gejala masukan (working memory) secara real-time dengan basis pengetahuan aturan (relasi rules pada tabel aturan) untuk ditarik kesimpulan menjadi diagnosa penyakit."),
        ("Visualisasi Proses Penelusuran Logika Dinamis (Solusi Rumusan Masalah 3): ", "Menyusun halaman diagnosa dan halaman detail konsultasi yang menyajikan analisis kecocokan untuk semua kemungkinan penyakit gigi. Penelusuran divisualisasikan dengan persentase bar, circle SVG, serta panel akordion interaktif yang menampilkan daftar gejala yang cocok (premis terpenuhi) dan gejala yang tidak cocok/kurang."),
        ("Penyajian Informasi Diagnosa Awal dan Laporan Cetak (Solusi Rumusan Masalah 4): ", "Menampilkan informasi deskripsi klinis penyakit gigi dan solusi rekomendasi penanganan medis awal secara detail. Halaman hasil diagnosa dan detail riwayat juga ditambahkan gaya cetak khusus (CSS print) sehingga dapat dicetak rapi ke kertas/PDF sebagai rujukan awal pemeriksaan ke drg. Hj. Rini Sutarti.")
    ]

    for title, desc in solutions:
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.6)
        p.paragraph_format.space_after = Pt(4)
        run_title = p.add_run(title)
        run_title.bold = True
        p.add_run(desc)


    # --- SECTION B ---
    p_subb = doc.add_paragraph('B. Pembahasan Algoritma', style='Skripsi Heading 2')
    p_subb.paragraph_format.space_before = Pt(18)
    p_subb.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "Metode Forward Chaining merupakan metode pencarian maju yang dimulai dari pengumpulan fakta-fakta klinis "
        "(dalam hal ini adalah gejala penyakit gigi yang dirasakan oleh pasien) untuk kemudian diuji dengan aturan-aturan "
        "yang ada di dalam basis pengetahuan (knowledge base). Logika penalaran inferensi ini secara umum dinyatakan sebagai berikut:"
    )

    # Formula Box style
    p_form = doc.add_paragraph()
    p_form.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_form.paragraph_format.space_before = Pt(6)
    p_form.paragraph_format.space_after = Pt(6)
    run_form = p_form.add_run("IF (Gejala_A AND Gejala_B AND ... Gejala_N) THEN (Penyakit_X)")
    run_form.bold = True
    run_form.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(12)
    p.add_run(
        "Untuk memberikan penjelasan yang konkrit mengenai proses inferensi, di bawah ini dipaparkan simulasi penelusuran "
        "Forward Chaining pada aplikasi Diagnosa Penyakit Gigi Metode Forward Chaining berdasarkan satu studi kasus pasien."
    )

    # 1. Studi Kasus dan Fakta Awal
    p_subb1 = doc.add_paragraph('1. Studi Kasus dan Fakta Awal (Working Memory)', style='Skripsi Heading 2')
    p_subb1.paragraph_format.left_indent = Inches(0.2)
    p_subb1.paragraph_format.space_before = Pt(6)
    p_subb1.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0.4)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "Misalkan seorang pasien melakukan konsultasi mandiri di Diagnosa Penyakit Gigi Metode Forward Chaining dan menginputkan keluhan berupa 4 gejala fisik. "
        "Gejala-gejala tersebut disimpan dalam Working Memory sistem sebagai fakta-fakta awal berikut:"
    )

    case_symptoms = [
        ("G003: ", "Gigi terlihat berlubang atau berwarna kehitaman"),
        ("G005: ", "Sensitivitas pada makanan/minuman dingin"),
        ("G006: ", "Sensitivitas pada makanan/minuman manis atau asam"),
        ("G025: ", "Nyeri hanya saat ada rangsangan (bukan spontan)")
    ]
    for code, desc in case_symptoms:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.6)
        p.paragraph_format.space_after = Pt(2)
        run_c = p.add_run(code)
        run_c.bold = True
        p.add_run(desc)

    # 2. Proses Pencocokan Aturan (Rules Matching)
    p_subb2 = doc.add_paragraph('2. Pencocokan Aturan dan Perhitungan Persentase', style='Skripsi Heading 2')
    p_subb2.paragraph_format.left_indent = Inches(0.2)
    p_subb2.paragraph_format.space_before = Pt(12)
    p_subb2.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0.4)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(12)
    p.add_run(
        "Mesin inferensi kemudian mencocokkan fakta awal di atas dengan aturan (rules) dari masing-masing 8 penyakit gigi. "
        "Rumus matematika yang digunakan untuk menghitung persentase kecocokan (K) adalah:"
    )

    p_math = doc.add_paragraph()
    p_math.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_math.paragraph_format.space_before = Pt(6)
    p_math.paragraph_format.space_after = Pt(6)
    run_math = p_math.add_run("K = (Jumlah Gejala Cocok / Total Gejala Aturan) x 100%")
    run_math.bold = True

    # Tracing table
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    hdr_cells = table.rows[0].cells
    headers = ['Kode', 'Nama Penyakit', 'Gejala Aturan', 'Gejala Cocok', 'Hasil Perhitungan', 'Persentase']
    widths = [Inches(0.6), Inches(1.8), Inches(1.3), Inches(1.1), Inches(1.3), Inches(0.9)]

    for idx, name in enumerate(headers):
        hdr_cells[idx].text = name
        hdr_cells[idx].width = widths[idx]
        set_cell_background(hdr_cells[idx], '0EA5E9')
        set_cell_margins(hdr_cells[idx], top=120, bottom=120, left=100, right=100)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(10)

    trace_data = [
        ('P001', 'Karies Gigi', 'G003, G005, G006, G023, G025', 'G003, G005, G006, G025', '(4 / 5) x 100%', '80.00%'),
        ('P002', 'Pulpitis', 'G001, G004, G005, G015, G024', 'G005', '(1 / 5) x 100%', '20.00%'),
        ('P003', 'Periodontitis', 'G007, G008, G009, G010, G011, G021, G022', '-', '0 / 7', '0.00%'),
        ('P004', 'Gingivitis', 'G007, G008, G010, G021', '-', '0 / 4', '0.00%'),
        ('P005', 'Abses Periapikal', 'G001, G002, G012, G013, G014, G015', '-', '0 / 6', '0.00%'),
        ('P006', 'Gigi Sensitif', 'G004, G005, G006, G022, G025', 'G005, G006, G025', '(3 / 5) x 100%', '60.00%'),
        ('P007', 'Sariawan', 'G016, G017, G018', '-', '0 / 3', '0.00%'),
        ('P008', 'Fraktur Gigi', 'G002, G004, G005, G019, G020', 'G005', '(1 / 5) x 100%', '20.00%'),
    ]

    for row_idx, data in enumerate(trace_data):
        row = table.add_row()
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.text = text
            cell.width = widths[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 4, 5] else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.size = Pt(9.5)
            # Alternating row colors
            if row_idx % 2 == 1:
                set_cell_background(cell, 'F8FAFC')

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0.4)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(12)
    p.add_run(
        "Berdasarkan tabel pencocokan di atas, sistem pakar mengurutkan penyakit dengan persentase kecocokan "
        "tertinggi sebagai diagnosa utama. Pada kasus ini, Karies Gigi (P001) menempati urutan pertama dengan "
        "persentase 80.00%, diikuti oleh Gigi Sensitif (P006) sebesar 60.00%, dan Pulpitis (P002) serta Fraktur Gigi (P008) "
        "masing-masing sebesar 20.00%. Hasil inilah yang kemudian disajikan kepada pasien di halaman hasil diagnosa "
        "dan disimpan dalam basis data riwayat."
    )


    # --- SECTION C ---
    p_subc = doc.add_paragraph('C. Pemodelan Perangkat Lunak', style='Skripsi Heading 2')
    p_subc.paragraph_format.space_before = Pt(18)
    p_subc.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "Aplikasi Diagnosa Penyakit Gigi Metode Forward Chaining dikembangkan menggunakan arsitektur web client-server. Server menggunakan bahasa pemrograman PHP "
        "yang bertugas memproses data masukan dan database MySQL untuk menyimpan data relational. "
        "Struktur tabel database yang dirancang untuk mendukung operasional sistem pakar ini meliputi:"
    )

    db_tables = [
        ("Tabel users: ", "Menyimpan data otentikasi admin dan dokter gigi (id, nama, username, password, role, created_at)."),
        ("Tabel penyakit: ", "Menyimpan kamus penyakit gigi (id, kode, nama, deskripsi, solusi, created_at)."),
        ("Tabel gejala: ", "Menyimpan kamus gejala klinis gigi dan mulut (id, kode, nama, created_at)."),
        ("Tabel aturan: ", "Menyimpan aturan logika relasional antara penyakit dan gejalanya (id, penyakit_id, gejala_id)."),
        ("Tabel konsultasi: ", "Menyimpan data hasil akhir konsultasi pasien (id, user_id, nama_pasien, tanggal, hasil_diagnosa, persentase)."),
        ("Tabel konsultasi_gejala: ", "Menyimpan detail gejala spesifik yang dipilih pasien pada suatu sesi konsultasi (id, konsultasi_id, gejala_id).")
    ]
    for title, desc in db_tables:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.6)
        p.paragraph_format.space_after = Pt(2)
        run_t = p.add_run(title)
        run_t.bold = True
        p.add_run(desc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(12)
    p.add_run(
        "Logika pemrograman engine Forward Chaining dibungkus ke dalam class PHP 'ForwardChaining' yang terletak di "
        "file forward_chaining.php. Ketika data dikirimkan dari form konsultasi, engine akan memuat seluruh daftar penyakit, "
        "melakukan query SQL berparameter (Prepared Statements) untuk mengambil relasi aturan gejala masing-masing penyakit, "
        "mencocokkannya dengan fakta input menggunakan fungsi in_array(), menghitung nilai persentase, mengurutkan hasil "
        "berdasarkan nilai persentase tertinggi menggunakan usort(), dan menyimpannya ke tabel konsultasi."
    )


    # --- SECTION D ---
    p_subd = doc.add_paragraph('D. Kelebihan dan Kelemahan Penelitian', style='Skripsi Heading 2')
    p_subd.paragraph_format.space_before = Pt(18)
    p_subd.paragraph_format.space_after = Pt(12)

    # 1. Kelebihan
    p_subd1 = doc.add_paragraph('1. Kelebihan Penelitian', style='Skripsi Heading 2')
    p_subd1.paragraph_format.left_indent = Inches(0.2)
    p_subd1.paragraph_format.space_before = Pt(6)
    p_subd1.paragraph_format.space_after = Pt(6)

    pros = [
        ("Efisiensi Waktu dan Biaya: ", "Memberikan diagnosa awal yang cepat (hitungan detik) secara gratis tanpa harus mengantre lama di klinik hanya untuk pemeriksaan dasar."),
        ("Antarmuka Modern dan Responsif: ", "Aplikasi dirancang dengan visual premium dark theme yang responsif, memudahkan diakses melalui perangkat komputer maupun smartphone secara nyaman."),
        ("Fitur Pencetakan Rujukan Mandiri: ", "Hasil diagnosa dapat dicetak langsung dalam tata letak dokumen formal yang rapi untuk dibawa saat konsultasi tatap muka."),
        ("Keamanan Pengelolaan Data: ", "Implementasi prepared statements pada query SQL meminimalkan celah kerentanan SQL Injection, dan password admin dienkripsi menggunakan algoritma bcrypt.")
    ]
    for title, desc in pros:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.6)
        p.paragraph_format.space_after = Pt(4)
        run_t = p.add_run(title)
        run_t.bold = True
        p.add_run(desc)

    # 2. Kelemahan
    p_subd2 = doc.add_paragraph('2. Kelemahan Penelitian', style='Skripsi Heading 2')
    p_subd2.paragraph_format.left_indent = Inches(0.2)
    p_subd2.paragraph_format.space_before = Pt(12)
    p_subd2.paragraph_format.space_after = Pt(6)

    cons = [
        ("Diagnosa Bersifat Screening Awal: ", "Sistem pakar tidak dapat menggantikan pemeriksaan fisik langsung oleh dokter gigi seperti rontgen gigi untuk mengetahui kedalaman lubang."),
        ("Keterbatasan Input Informasi Medis: ", "Penentuan diagnosa saat ini murni hanya berdasarkan checkbox keluhan gejala, tanpa didukung input multimedia (seperti foto visual kondisi gigi pasien)."),
        ("Ketergantungan Terhadap Kejujuran Pasien: ", "Jika pasien salah memilih gejala atau tidak mengenali gejala yang mereka alami secara akurat, sistem akan menghasilkan persentase diagnosa yang kurang tepat.")
    ]
    for title, desc in cons:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.6)
        p.paragraph_format.space_after = Pt(4)
        run_t = p.add_run(title)
        run_t.bold = True
        p.add_run(desc)


    # -------------------------------------------------------------
    # BAB V PENUTUP
    # -------------------------------------------------------------
    doc.add_page_break()

    p_bab5 = doc.add_paragraph('BAB V', style='Skripsi Heading 1')
    p_bab5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bab5.paragraph_format.space_before = Pt(12)
    p_bab5.paragraph_format.space_after = Pt(6)

    p_title5 = doc.add_paragraph('PENUTUP', style='Skripsi Heading 1')
    p_title5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title5.paragraph_format.space_after = Pt(24)

    # --- SECTION A ---
    p_suba5 = doc.add_paragraph('A. Kesimpulan', style='Skripsi Heading 2')
    p_suba5.paragraph_format.space_before = Pt(12)
    p_suba5.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "Berdasarkan hasil analisis, perancangan, implementasi, dan pengujian yang telah dilakukan pada Sistem Pakar "
        "Penyakit Gigi (Diagnosa Penyakit Gigi Metode Forward Chaining) menggunakan metode Forward Chaining di Praktik Mandiri Drg. Hj. Rini Sutarti, "
        "dapat ditarik beberapa kesimpulan sebagai berikut:"
    )

    conclusions = [
        "Aplikasi Diagnosa Penyakit Gigi Metode Forward Chaining berhasil dikembangkan sebagai platform berbasis web yang memudahkan pasien dalam melakukan diagnosa awal mandiri penyakit gigi berdasarkan 25 keluhan gejala klinis.",
        "Metode Forward Chaining terbukti sangat cocok diimplementasikan dalam sistem pakar ini untuk melakukan penelusuran maju dari sekumpulan fakta gejala klinis menuju kesimpulan diagnosa penyakit gigi dengan akurasi logika yang terstruktur sesuai basis pengetahuan pakar.",
        "Sistem penghitungan persentase kecocokan memberikan gambaran tingkat kemiripan klinis yang objektif bagi pasien untuk memahami estimasi diagnosa penyakit.",
        "Halaman Riwayat Konsultasi dan fungsi pengelolaan panel admin berhasil mendokumentasikan rekam data keluhan pasien secara teratur, meminimalkan administrasi manual di klinik."
    ]
    for idx, c in enumerate(conclusions):
        p = doc.add_paragraph(style='List Number')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(c)

    # --- SECTION B ---
    p_subb5 = doc.add_paragraph('B. Saran', style='Skripsi Heading 2')
    p_subb5.paragraph_format.space_before = Pt(18)
    p_subb5.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "Untuk pengembangan aplikasi Diagnosa Penyakit Gigi Metode Forward Chaining lebih lanjut agar memiliki manfaat medis yang lebih luas di masa depan, "
        "penulis memberikan beberapa saran konstruktif sebagai berikut:"
    )

    suggestions = [
        "Pengembangan modul diagnosa visual menggunakan teknologi Computer Vision untuk mendeteksi karies atau karang gigi secara otomatis melalui unggahan foto kondisi mulut pasien.",
        "Pengembangan fitur konsultasi lanjutan berbasis real-time chat (telemedicine) yang menghubungkan pasien langsung dengan Drg. Hj. Rini Sutarti setelah hasil screening awal keluar.",
        "Integrasi modul sistem penjadwalan janji temu online (appointment scheduling) ke dalam aplikasi Diagnosa Penyakit Gigi Metode Forward Chaining agar pasien yang membutuhkan tindakan fisik dapat langsung memesan jadwal pemeriksaan di klinik."
    ]
    for idx, s in enumerate(suggestions):
        p = doc.add_paragraph(style='List Number')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(s)


    # Save document
    filename = 'Hasil_dan_Pembahasan_Diagnosa_Penyakit_Gigi_Metode_Forward_Chaining.docx'
    doc.save(filename)
    print(f"Document saved successfully as {filename}")

if __name__ == '__main__':
    create_document()
